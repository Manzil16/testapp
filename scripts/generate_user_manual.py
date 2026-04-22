"""Generate VehicleGrid end-user manual as a DOCX."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x00, 0xBF, 0xA5)
DARK = RGBColor(0x10, 0x1A, 0x2E)
GREY = RGBColor(0x4A, 0x55, 0x68)
LIGHT_GREY = RGBColor(0x8A, 0x93, 0xA6)

doc = Document()

# ---------- Page margins ----------
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

# ---------- Default styles ----------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = DARK

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def add_heading(text, level=1, color=TEAL, size=None):
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes.get(level, 12))
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p

def add_para(text, *, bold=False, italic=False, color=None, size=11, align=None,
             space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or DARK
    run.font.name = "Calibri"
    return p

def add_bullet(text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + indent * 0.6)
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run()
    if p.runs:
        p.runs[0].text = text
    else:
        run.text = text
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        r.font.name = "Calibri"
    return p

def add_numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        r.font.name = "Calibri"
    return p

def add_divider():
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "00BFA5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def add_callout(title, body, fill="E6FAF5", border="00BFA5"):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, fill)
    # border color
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:color"), border)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

    # title
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    r = p_title.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = TEAL
    r.font.name = "Calibri"

    # body
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(0)
    rb = p_body.add_run(body)
    rb.font.size = Pt(10.5)
    rb.font.color.rgb = DARK
    rb.font.name = "Calibri"

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)

def add_step_table(steps):
    """steps = list of (step_no, title, description)"""
    table = doc.add_table(rows=len(steps), cols=2)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(1.4)
        row.cells[1].width = Cm(14.5)
    for i, (num, title, desc) in enumerate(steps):
        num_cell = table.rows[i].cells[0]
        body_cell = table.rows[i].cells[1]
        set_cell_bg(num_cell, "00BFA5")
        num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p = num_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(num))
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Calibri"

        tp = body_cell.paragraphs[0]
        tp.paragraph_format.space_after = Pt(2)
        tr = tp.add_run(title)
        tr.bold = True
        tr.font.size = Pt(11.5)
        tr.font.color.rgb = DARK
        tr.font.name = "Calibri"

        if desc:
            dp = body_cell.add_paragraph()
            dp.paragraph_format.space_after = Pt(0)
            dr = dp.add_run(desc)
            dr.font.size = Pt(10.5)
            dr.font.color.rgb = GREY
            dr.font.name = "Calibri"

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

def page_break():
    doc.add_page_break()

# ---------- Footer with team names ----------
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Manzil  •  Aman  •  Jaydan  •  Tausif")
fr.font.size = Pt(10)
fr.font.color.rgb = LIGHT_GREY
fr.font.name = "Calibri"
fr.bold = True

# =================================================================
# COVER PAGE
# =================================================================
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(140)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run("VehicleGrid")
run.bold = True
run.font.size = Pt(54)
run.font.color.rgb = TEAL
run.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Peer-to-Peer EV Charging Marketplace")
sr.font.size = Pt(16)
sr.font.color.rgb = DARK
sr.italic = True
sr.font.name = "Calibri"

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tagline.add_run("Find, book, and power up — or earn from your home charger.")
tr.font.size = Pt(12)
tr.font.color.rgb = GREY
tr.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(80)

manual_label = doc.add_paragraph()
manual_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
ml = manual_label.add_run("END-USER MANUAL")
ml.bold = True
ml.font.size = Pt(14)
ml.font.color.rgb = TEAL
ml.font.name = "Calibri"

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
vr = version.add_run("Version 1.0  |  April 2026")
vr.font.size = Pt(11)
vr.font.color.rgb = GREY
vr.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(120)

team = doc.add_paragraph()
team.alignment = WD_ALIGN_PARAGRAPH.CENTER
tm = team.add_run("Prepared by")
tm.font.size = Pt(10)
tm.font.color.rgb = LIGHT_GREY
tm.font.name = "Calibri"

names = doc.add_paragraph()
names.alignment = WD_ALIGN_PARAGRAPH.CENTER
nm = names.add_run("Manzil  •  Aman  •  Jaydan  •  Tausif")
nm.bold = True
nm.font.size = Pt(12)
nm.font.color.rgb = DARK
nm.font.name = "Calibri"

page_break()

# =================================================================
# TABLE OF CONTENTS
# =================================================================
add_heading("Contents", level=1)
add_divider()

toc_items = [
    ("1.", "Introduction", "3"),
    ("2.", "Installation and Access", "4"),
    ("3.", "Getting Started", "5"),
    ("4.", "Using VehicleGrid as a Driver", "7"),
    ("5.", "Using VehicleGrid as a Host", "11"),
    ("6.", "Administrator Tools", "14"),
    ("7.", "Payments and Refunds", "15"),
    ("8.", "Troubleshooting", "16"),
    ("9.", "Contact and Support", "17"),
]
toc_table = doc.add_table(rows=len(toc_items), cols=3)
toc_table.autofit = False
for i, (num, label, pg) in enumerate(toc_items):
    row = toc_table.rows[i]
    row.cells[0].width = Cm(1.0)
    row.cells[1].width = Cm(13.5)
    row.cells[2].width = Cm(1.5)

    n = row.cells[0].paragraphs[0]
    nr = n.add_run(num)
    nr.bold = True; nr.font.size = Pt(11); nr.font.color.rgb = TEAL; nr.font.name = "Calibri"

    t = row.cells[1].paragraphs[0]
    tr = t.add_run(label)
    tr.font.size = Pt(11); tr.font.color.rgb = DARK; tr.font.name = "Calibri"

    pp = row.cells[2].paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr = pp.add_run(pg)
    pr.font.size = Pt(11); pr.font.color.rgb = GREY; pr.font.name = "Calibri"

page_break()

# =================================================================
# 1. INTRODUCTION
# =================================================================
add_heading("1. Introduction", level=1)
add_divider()

add_para(
    "Welcome to VehicleGrid — a friendly mobile app that helps electric-vehicle (EV) "
    "drivers find places to charge, and helps home-charger owners earn money by sharing "
    "their charger with the community. Think of it as Airbnb, but for EV chargers.",
    space_after=8,
)
add_para(
    "This manual is written for everyday users. You do not need any technical background. "
    "Follow along page by page, or jump to the section that matches what you want to do.",
    space_after=12,
)

add_heading("Who is VehicleGrid for?", level=2)
add_bullet("Drivers — people with an electric vehicle who want to book a charger near home, work, or along a road trip.")
add_bullet("Hosts — people who own a home charger and want to earn by renting it out by the hour or by the kilowatt-hour (kWh).")
add_bullet("Administrators — the VehicleGrid team, who review new chargers and keep the platform safe.")

add_heading("Key Features at a Glance", level=2)
features = [
    ("Interactive map and list", "Browse chargers around you, filter by plug type and price, and tap any pin to see details."),
    ("Smart trip planner", "Enter where you are going and let VehicleGrid suggest chargers along the way that add the smallest detour."),
    ("Transparent pricing", "See the price per kWh, an estimate for your session, and the 10% platform fee before you confirm anything."),
    ("Live booking tracker", "Watch your booking move from Requested → Approved → Charging → Complete, with a live countdown and actual energy used."),
    ("Safe payments", "Your card is only held, not charged, until your session ends. You only pay for the kWh you actually use."),
    ("Host dashboard", "List a charger, set your price and hours, and track your earnings and upcoming bookings."),
    ("Reviews and trust", "Read ratings and comments from other drivers, and see how quickly a host usually responds."),
    ("Digital verification", "Drivers and hosts complete a short digital check (email, phone, ID, card) so everyone on the platform is trusted."),
]
for title, desc in features:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    rt = p.add_run(title + " — ")
    rt.bold = True; rt.font.size = Pt(11); rt.font.color.rgb = TEAL; rt.font.name = "Calibri"
    rd = p.add_run(desc)
    rd.font.size = Pt(11); rd.font.color.rgb = DARK; rd.font.name = "Calibri"

add_callout(
    "A quick word about safety",
    "VehicleGrid never charges your card without your approval. We place a temporary "
    "hold when you book, then capture the real amount based on the actual energy you used. "
    "If you charge less than estimated, you pay less — automatically.",
)

page_break()

# =================================================================
# 2. INSTALLATION AND ACCESS
# =================================================================
add_heading("2. Installation and Access", level=1)
add_divider()

add_para(
    "VehicleGrid runs on iPhone and Android. You can start using the app in under two minutes.",
    space_after=8,
)

add_heading("What you need", level=2)
add_bullet("A smartphone running iOS 15 or later, or Android 10 or later.")
add_bullet("An internet connection (Wi-Fi or mobile data).")
add_bullet("An email address and a mobile phone number.")
add_bullet("(Drivers) a debit or credit card for payment.")
add_bullet("(Hosts) a bank account and a photo ID for payouts and verification.")

add_heading("Installing the app", level=2)
add_step_table([
    (1, "Open the app store on your phone",
     "Apple App Store on iPhone or Google Play Store on Android."),
    (2, "Search for \"VehicleGrid\"",
     "Look for the teal VehicleGrid icon with the white charging plug symbol."),
    (3, "Tap Install (or Get)",
     "The app will download and appear on your home screen. No extra setup is required."),
    (4, "Tap the VehicleGrid icon to open",
     "You will land on the welcome screen where you can sign in or create an account."),
])

add_callout(
    "Using a demo build?",
    "If your teacher or reviewer gave you a demo link, open it in Expo Go (installed "
    "from the app store) and scan the QR code. The VehicleGrid logo will appear inside "
    "Expo Go's project list.",
    fill="FFF8E6", border="E0A800",
)

page_break()

# =================================================================
# 3. GETTING STARTED
# =================================================================
add_heading("3. Getting Started", level=1)
add_divider()

add_heading("3.1 Creating your account", level=2)
add_step_table([
    (1, "Tap \"Create an account\" on the welcome screen",
     "You will see three role cards: Driver, Host, and Admin (staff only)."),
    (2, "Pick the role that matches what you want to do",
     "You can add another role later from Settings — you are not locked in."),
    (3, "Continue with Google, or enter an email and password",
     "If you choose email, pick a password with at least 8 characters."),
    (4, "Enter your full name and mobile number",
     "Use a real number — we will send a 6-digit code to verify it."),
    (5, "Type the 6-digit verification code",
     "The code arrives by SMS and is valid for 10 minutes. Tap Verify."),
    (6, "Agree to the Terms and tap Finish",
     "You are now signed in and can start using the app."),
])

add_heading("3.2 Signing in on another device", level=2)
add_numbered("Open the app and tap \"Sign in\" on the welcome screen.")
add_numbered("Enter the email and password you used, or tap \"Continue with Google\".")
add_numbered("If prompted, re-verify your phone with a fresh 6-digit code.")

add_heading("3.3 The main screen layout", level=2)
add_para(
    "Once you are signed in, you will see a bar of tabs along the bottom of the screen. "
    "Tabs change depending on your role:",
    space_after=6,
)

tab_table = doc.add_table(rows=5, cols=3)
tab_table.style = "Light Grid Accent 1"
headers = ["Tab", "Who sees it", "What it does"]
for i, h in enumerate(headers):
    cell = tab_table.rows[0].cells[i]
    set_cell_bg(cell, "00BFA5")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.name = "Calibri"

tab_rows = [
    ("Discover", "Drivers", "Map and list of chargers near you."),
    ("Trips", "Drivers", "Plan a journey and find chargers along the way."),
    ("Bookings", "Drivers & Hosts", "All current and past sessions."),
    ("Profile", "Everyone", "Your account, vehicle, payment, and settings."),
]
for i, (a, b, c) in enumerate(tab_rows, start=1):
    cells = tab_table.rows[i].cells
    for j, val in enumerate((a, b, c)):
        cell = cells[j]
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10.5); r.font.color.rgb = DARK; r.font.name = "Calibri"
        if j == 0:
            r.bold = True

add_para(" ", space_after=6)

add_heading("3.4 Completing your profile", level=2)
add_para(
    "Before you can book or list a charger, you need to complete a short digital check. "
    "Open Profile and tap \"Verification\". A progress card shows what is still needed.",
    space_after=6,
)

add_para("Drivers must complete:", bold=True, space_after=2)
add_bullet("Email verified (automatic when you sign up).")
add_bullet("Phone verified (the 6-digit code you entered above).")
add_bullet("Payment method added — tap \"Add card\" and enter your card details on Stripe's secure form.")

add_para("Hosts must complete everything above, plus:", bold=True, space_before=6, space_after=2)
add_bullet("Upload a photo of a government-issued ID (driver licence or passport).")
add_bullet("Connect a payout account through Stripe Connect — add your bank and tax details.")

add_callout(
    "Your data is safe",
    "Card numbers are handled by Stripe, not VehicleGrid. Your ID photo is stored in a "
    "private, encrypted bucket and is only visible to the admin reviewing your profile.",
)

page_break()

# =================================================================
# 4. DRIVER FEATURES
# =================================================================
add_heading("4. Using VehicleGrid as a Driver", level=1)
add_divider()

add_heading("4.1 Finding a charger near you", level=2)
add_step_table([
    (1, "Tap the Discover tab",
     "The app will ask for location permission the first time — tap Allow."),
    (2, "Browse the map",
     "Teal pins are available chargers. Pinch to zoom or drag the map. A small preview card "
     "appears at the bottom when you tap a pin."),
    (3, "Switch to list view if you prefer",
     "Tap the list icon at the top right. Chargers appear sorted by distance from you."),
    (4, "Use filters to narrow the results",
     "Filter by plug type (Type 2, CCS2, CHAdeMO), minimum power (22 kW+, 50 kW+), or price."),
    (5, "Tap \"View details →\" on any preview",
     "The full charger page opens with photos, specs, reviews, and a Book button."),
])

add_heading("4.2 Reading a charger's detail page", level=2)
add_para("Scroll down the page — each section answers a common question:", space_after=4)
add_bullet("Hero photo — is this the right spot?")
add_bullet("Star rating and session count — can I trust this charger?")
add_bullet("Today's availability bar — is it free when I need it? (Teal = free, red = booked, grey = past)")
add_bullet("Specs — will it fit my car? (Plug type, max power, cable length)")
add_bullet("Amenities — what is nearby? (Cafe, Wi-Fi, covered parking)")
add_bullet("Host profile — who am I dealing with?")
add_bullet("Recent reviews — what do other drivers say?")
add_para(
    "The price bar stays pinned at the bottom of the screen as you scroll, so the Book button is always one tap away.",
    italic=True, color=GREY, space_before=4, space_after=6,
)

add_heading("4.3 Booking a session", level=2)
add_step_table([
    (1, "Tap \"Book this charger\" on the detail page",
     "You move to the Confirm Booking screen."),
    (2, "Pick a start date and time",
     "Tap the date chip to open the calendar, then the time chip to set the hour and minute."),
    (3, "Adjust the battery target slider",
     "Drag the slider to the percentage you want your car charged to (for example, from 30% to 80%)."),
    (4, "Review the live estimate",
     "VehicleGrid shows the estimated energy (kWh), how long it will take, and the total cost "
     "including the 10% platform fee."),
    (5, "Tap \"Request booking\"",
     "Your card is placed on hold for the estimated total — not charged yet. The booking appears "
     "in your Bookings tab as \"Waiting for host\"."),
])

add_callout(
    "Not sure about the numbers?",
    "The estimate uses your vehicle's battery size, the charger's power, and a 10% efficiency "
    "loss (heat lost in the cable). Your final bill is based on actual kWh, so you never pay "
    "for more energy than your car received.",
)

add_heading("4.4 Tracking a live booking", level=2)
add_para(
    "Open the Bookings tab and tap the booking you want to track. The Booking Details page "
    "is the most important screen for you as a driver — it updates in real time.",
    space_after=6,
)

add_para("What each status means:", bold=True, space_after=2)
status_table = doc.add_table(rows=7, cols=2)
status_rows = [
    ("Requested", "The host has up to 24 hours to approve. Your card is on hold."),
    ("Approved", "You're confirmed. A live countdown shows time until session start."),
    ("Charging", "Session in progress. A pulse dot and running kWh readout update live."),
    ("Completed", "Your final receipt appears with the actual kWh and amount charged."),
    ("Cancelled", "Either party cancelled. The refund matrix decides what you get back."),
    ("Missed", "You did not arrive within the 15-minute grace window. Hold is released but no refund."),
    ("Expired / Declined", "Host did not respond in 24 hours, or declined. Hold is released immediately."),
]
for i, (a, b) in enumerate(status_rows):
    cells = status_table.rows[i].cells
    cells[0].width = Cm(4.0)
    set_cell_bg(cells[0], "E6FAF5")
    p = cells[0].paragraphs[0]
    r = p.add_run(a)
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = TEAL; r.font.name = "Calibri"

    p2 = cells[1].paragraphs[0]
    r2 = p2.add_run(b)
    r2.font.size = Pt(10.5); r2.font.color.rgb = DARK; r2.font.name = "Calibri"

add_para(" ", space_after=6)

add_heading("4.5 Arriving at the charger", level=2)
add_numbered("Use the Directions button to open Apple Maps or Google Maps.")
add_numbered("When you plug in, tap \"I've arrived\" on the Booking Details page. This stops the 15-minute no-show timer.")
add_numbered("The status becomes Charging and you can watch the live kWh readout.")
add_numbered("When your car is full (or you need to leave), tap \"End session\". The final amount is captured and the hold is released.")

add_heading("4.6 Planning a long trip", level=2)
add_step_table([
    (1, "Tap the Trips tab",
     "The trip planner opens."),
    (2, "Enter your starting point and destination",
     "You can type an address or tap \"Use my location\" for the start."),
    (3, "Tap \"Find chargers along the way\"",
     "VehicleGrid draws your route and highlights chargers that add less than 15% to your trip."),
    (4, "Review the ranked list",
     "Chargers are scored on detour, price, power, and availability. Tap any one to see its full detail page."),
    (5, "Book the one that suits you",
     "Booking works exactly like the single-charger flow above."),
])

add_heading("4.7 Cancelling a booking", level=2)
add_numbered("Open the booking in the Bookings tab.")
add_numbered("Tap \"Cancel booking\" at the bottom of the page.")
add_numbered("Pick a reason from the list (changed plans, found closer charger, and so on).")
add_numbered("Confirm. Your refund is calculated automatically — see Section 7 for the full refund matrix.")

page_break()

# =================================================================
# 5. HOST FEATURES
# =================================================================
add_heading("5. Using VehicleGrid as a Host", level=1)
add_divider()

add_para(
    "Hosts earn money by sharing their home charger. This section walks you through listing "
    "a charger, managing bookings, and tracking earnings.",
    space_after=8,
)

add_heading("5.1 Listing your first charger", level=2)
add_step_table([
    (1, "Tap the Host Home tab, then \"Add a charger\"",
     "Available after you finish host verification in Section 3.4."),
    (2, "Take at least two clear photos",
     "Include a wide shot and a close-up of the plug. More photos build trust and raise your ranking."),
    (3, "Enter the name and address",
     "Use something descriptive like \"Smith Family Garage — 7 kW Type 2\". The app will geocode the address automatically."),
    (4, "Select your plug type and maximum power",
     "Pick from Type 2, CCS2, or CHAdeMO. Enter the maximum kW your charger delivers."),
    (5, "Set the price per kWh",
     "Must be between $0.20 and $2.50. A helper shows the market average for your suburb."),
    (6, "Choose your available hours",
     "Pick the days and time windows your charger can be booked. You can change these later."),
    (7, "Add amenities (optional)",
     "Tick any nearby perks: Wi-Fi, cafe, covered parking, toilet, well-lit area."),
    (8, "Tap \"Submit for review\"",
     "An admin reviews your listing within 24 hours. You will get a push notification when it is approved."),
])

add_callout(
    "How are listings approved?",
    "Admins score each listing out of 100 across five criteria: photos, specs, location accuracy, "
    "access, and pricing. A score of 85 or more is approved automatically. Anything below 45 is rejected. "
    "In between, an admin will add a short note explaining what to fix.",
)

add_heading("5.2 Managing booking requests", level=2)
add_step_table([
    (1, "Open the Host Bookings tab",
     "New requests appear at the top with an amber dot."),
    (2, "Tap the request to see full details",
     "Driver name, vehicle, requested time window, estimated kWh, and the amount on hold."),
    (3, "Approve or Decline",
     "You have 24 hours to respond. Approving maintains the card hold; declining releases it immediately."),
    (4, "The driver receives a push notification",
     "They will also see the updated status in their Bookings tab."),
])

add_para(
    "Tip: Respond quickly. Your average response time is shown on your charger's detail page "
    "(for example, \"Usually responds within 2 hours\"). Faster responders rank higher in search.",
    italic=True, color=GREY, space_after=6,
)

add_heading("5.3 During and after a session", level=2)
add_numbered("When the driver taps \"I've arrived\", the booking status changes to Charging.")
add_numbered("You do not need to do anything during the session — the charger reports the energy used automatically.")
add_numbered("When the session ends, VehicleGrid captures the final amount from the driver and credits your payout (their subtotal minus the 10% platform fee).")
add_numbered("Payouts land in your linked bank account within 2–3 business days.")

add_heading("5.4 Tracking earnings", level=2)
add_step_table([
    (1, "Open the Host Home tab",
     "The Earnings card shows this week's total, last week's total, and your lifetime revenue."),
    (2, "Tap \"View analytics\"",
     "See a breakdown per charger: sessions completed, average price, busiest day of the week, and total kWh delivered."),
    (3, "Download a monthly statement",
     "From the Analytics page, tap \"Download CSV\" — useful for tax time."),
])

add_heading("5.5 Editing or pausing a charger", level=2)
add_bullet("To change price or hours — open Host Chargers, tap the charger, tap Edit, and save.")
add_bullet("To temporarily pause — open the charger and toggle \"Pause listing\". No new bookings can be made until you unpause.")
add_bullet("To delete a charger — tap Edit → Delete. All upcoming bookings will be cancelled with full refunds.")

page_break()

# =================================================================
# 6. ADMIN
# =================================================================
add_heading("6. Administrator Tools", level=1)
add_divider()

add_para(
    "The admin tabs are only visible to VehicleGrid staff. If you are a regular user, you can "
    "skip this section.",
    italic=True, color=GREY, space_after=8,
)

add_heading("6.1 Platform Overview", level=2)
add_para("The main admin screen is a searchable feed of everything happening on VehicleGrid.", space_after=4)
add_bullet("Type any name, email, charger, or booking ID in the search bar.")
add_bullet("Use the filter chips to see only Needs Action, Payments, Cancellations, New Chargers, or Completed Sessions.")
add_bullet("Tap any row to expand full details and take action (refund, suspend, approve, and so on).")
add_bullet("The stat row at the top shows today's revenue, active sessions, pending approvals, and new users.")

add_heading("6.2 Charger Verification Queue", level=2)
add_numbered("Open the Verify tab. Pending chargers are listed newest first.")
add_numbered("Tap a charger to view every submitted photo full-screen.")
add_numbered("Score the listing across five rubric dimensions (each out of 20).")
add_numbered("Approve, Reject, or Request Changes. Your action is written to the event log instantly.")

add_heading("6.3 Trust and Disputes", level=2)
add_bullet("Repeat cancellers — drivers with more than two cancellations. Warn or suspend as needed.")
add_bullet("Slow hosts — hosts whose average response time exceeds 2 hours. Send a nudge or demote them in ranking.")
add_bullet("Disputed sessions — bookings cancelled with a reason given. Refund or escalate as appropriate.")

page_break()

# =================================================================
# 7. PAYMENTS AND REFUNDS
# =================================================================
add_heading("7. Payments and Refunds", level=1)
add_divider()

add_heading("7.1 How you are charged", level=2)
add_para(
    "VehicleGrid uses a \"hold and reconcile\" model. When you book, your card is held for the "
    "estimated total but not charged. When your session ends, we charge only for the real kWh "
    "used, and release the rest of the hold.",
    space_after=6,
)

add_para("Example:", bold=True, space_after=2)
add_bullet("Estimated 35 kWh × $0.55 per kWh = $19.25 subtotal")
add_bullet("Plus 10% platform fee = $21.18 held on your card")
add_bullet("Actual session used 31.6 kWh. Final charge = $17.38 × 1.10 = $19.12")
add_bullet("The remaining $2.06 is released back automatically")

add_heading("7.2 Refund policy", level=2)
refund = doc.add_table(rows=7, cols=2)
refund.style = "Light Grid Accent 1"
refund_rows = [
    ("Situation", "Refund"),
    ("Driver cancels more than 2 hours before start", "100%"),
    ("Driver cancels less than 2 hours before start", "50%"),
    ("Driver misses booking (grace window lapses)", "0%"),
    ("Host cancels after approval", "100%"),
    ("Charger fails mid-session", "Proportional — only actual kWh is charged"),
    ("Host declines before approving", "100% — hold released immediately"),
]
for i, (a, b) in enumerate(refund_rows):
    cells = refund.rows[i].cells
    if i == 0:
        set_cell_bg(cells[0], "00BFA5")
        set_cell_bg(cells[1], "00BFA5")
        for c in cells:
            p = c.paragraphs[0]
            r = p.runs[0] if p.runs else p.add_run()
            r.text = a if c is cells[0] else b
            r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.name = "Calibri"
    else:
        p1 = cells[0].paragraphs[0]; r1 = p1.add_run(a)
        r1.font.size = Pt(10.5); r1.font.color.rgb = DARK; r1.font.name = "Calibri"
        p2 = cells[1].paragraphs[0]; r2 = p2.add_run(b)
        r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = TEAL; r2.font.name = "Calibri"

add_para(" ", space_after=6)

add_heading("7.3 Host payouts", level=2)
add_bullet("You receive 90% of the subtotal (the remaining 10% is the platform fee).")
add_bullet("Payouts are transferred through Stripe Connect to your linked bank account.")
add_bullet("Processing time is 2–3 business days after a session completes.")
add_bullet("You can track pending and paid payouts in Host Home → Earnings.")

page_break()

# =================================================================
# 8. TROUBLESHOOTING
# =================================================================
add_heading("8. Troubleshooting", level=1)
add_divider()

tr_items = [
    ("I didn't receive my phone verification code",
     "Wait 60 seconds and tap \"Resend code\". Check that the country code in your number is correct. "
     "If it still doesn't arrive, sign out and sign back in to trigger a fresh send."),
    ("The map is empty and shows no chargers",
     "Pull down to refresh. Make sure you have given the app location permission "
     "(iPhone Settings → VehicleGrid → Location → While Using). Zoom out — chargers may be more than 5 km away."),
    ("My card was declined when adding a payment method",
     "This is a Stripe check, not a real charge. Try a different card, or contact your bank to allow "
     "online payments. Prepaid and some international cards may not work."),
    ("The host isn't responding to my request",
     "Requests expire after 24 hours and the hold is released automatically. In the meantime you can "
     "cancel the request at any time for a full refund."),
    ("I tapped \"I've arrived\" but the status didn't change",
     "Check your internet connection and pull down to refresh. If the charger shows \"Charging\" but "
     "your car isn't receiving power, contact the host using the chat button."),
    ("My final charge looks higher than the estimate",
     "This happens when your car used more energy than planned (for example, battery warming in cold "
     "weather). Your receipt shows the exact kWh delivered. If the number looks wrong, tap \"Report issue\" "
     "on the booking."),
    ("I can't submit my charger listing",
     "The Submit button is only enabled when you have added at least two photos, a price within the "
     "$0.20–$2.50 range, and the address matches the GPS location within 500 metres."),
    ("I never got a push notification for my booking",
     "Go to your phone Settings → Notifications → VehicleGrid and make sure notifications are enabled. "
     "You also need to be signed in at least once on this device to register."),
]
for q, a in tr_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Q.  " + q)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK; r.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Cm(0.6)
    r2 = p2.add_run("A.  " + a)
    r2.font.size = Pt(10.5); r2.font.color.rgb = GREY; r2.font.name = "Calibri"

page_break()

# =================================================================
# 9. CONTACT AND SUPPORT
# =================================================================
add_heading("9. Contact and Support", level=1)
add_divider()

add_para(
    "We want your VehicleGrid experience to be smooth. If something isn't working or "
    "you have a suggestion, reach out — we read every message.",
    space_after=10,
)

contact = doc.add_table(rows=4, cols=2)
contact.autofit = False
contact_rows = [
    ("In-app Help Centre", "Profile → Help — browse common articles and open a support ticket."),
    ("Email", "support@vehiclegrid.app — replies within one business day."),
    ("Trust and Safety", "trust@vehiclegrid.app — use this for urgent safety or dispute issues."),
    ("Team (this project)", "Manzil, Aman, Jaydan, Tausif"),
]
for i, (a, b) in enumerate(contact_rows):
    cells = contact.rows[i].cells
    cells[0].width = Cm(5.5)
    cells[1].width = Cm(11.0)
    set_cell_bg(cells[0], "E6FAF5")
    p = cells[0].paragraphs[0]; r = p.add_run(a)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = TEAL; r.font.name = "Calibri"
    p2 = cells[1].paragraphs[0]; r2 = p2.add_run(b)
    r2.font.size = Pt(11); r2.font.color.rgb = DARK; r2.font.name = "Calibri"

add_para(" ", space_after=10)

add_callout(
    "Thank you for using VehicleGrid",
    "Every booking powers a cleaner commute and supports someone in your community. "
    "Drive green, charge smart, and safe travels.",
)

# =================================================================
# Save
# =================================================================
out = "/Users/manzildahal/Desktop/testapp/VehicleGrid_User_Manual.docx"
doc.save(out)
print("Saved:", out)
